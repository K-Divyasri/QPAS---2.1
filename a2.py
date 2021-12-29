from a1 import *

if (status == True):
    from re import search
    import docx
    import os
    # from guichanged import *
    import xlwt
    # import GUI

    import openpyxl as xl

    from xlwt import Workbook
    import win32com.client as win32
    from docx import Document

    AddComments = []
    details = []  # Add Add comments
    sub_code = ""
    outer_levels = []

    def obtain_details(row):
        global details
        temp = []
        for i in row:
            if i not in temp:
                temp.append(i)

        details.append(temp)


    def table_text(fp):
        document = docx.Document(fp)
        fultab = []
        qno = 1
        or_qnos = []
        msub = {}
        temp = []
        ortype = ['(Or)', '(OR)', 'OR', 'Or']
        for tno in range(0,len(document.tables)):  # Going through each table and adding it to the list $$$$$ MAKE end value of range NOO (4 for CAT1 doc)
            table = document.tables[tno]
            row_data = {}
            for row in table.rows:
                flag = 0
                texts = (cell.text for cell in row.cells)
                row_data[qno] = [text for text in texts]
                if flag == 0:
                    for every in ortype:
                        if flag == 0:
                            temp = []
                            for ea in row_data[qno]:
                                temp.append(ea.strip())
                            if (every == temp or [every] * 3 == temp or every in temp) and flag == 0:
                                or_qnos.append([len(fultab) + 1, qno - 1, qno])
                                row_data.pop(qno)
                                qno -= 1
                                flag = -1
                    if (row_data[qno] == ['', '', ''] or row_data[qno] == [] or row_data[qno] == ['']) and flag == 0:
                        row_data.pop(qno)
                        qno -= 1
                        flag = -1
                    elif (len(row_data[qno]) != 3 or ("K" not in row_data[qno][0] or "CO" not in row_data[qno][
                        len(row_data[qno]) - 1])) and flag == 0:
                        obtain_details(row_data[qno])
                        row_data.pop(qno)
                        qno -= 1
                        flag = -1
                if flag == 0:
                    if row_data[qno][0].count('K') > 1 or row_data[qno][1].casefold().count(
                            ' marks'.casefold()) > 1 or len(sub_marks(row_data[qno][1])) > 1:  # or ('i' in row_data[qno][1] and 'ii' in row_data[qno][1]) or 'iii' in row_data[qno][1]:
                        row_data[qno], msub[qno] = subquestions(row_data[qno]), sub_marks(row_data[qno][1])
                    else:
                        row_data[qno][0] = kfinder(row_data[qno][0])
                        row_data[qno][2] = cofinder(row_data[qno][2])

                qno += 1
            if row_data != {}:
                fultab.append(row_data)

        det_xl(details)
        return fultab, or_qnos, msub  # Returning the full table, or question numbers, marks of the subquestions


    def det_xl(details):
        college = "SSN COLLEGE OF ENGINEERING"
        department = ""
        exam = ""
        date = ""
        subject = ""
        branch = ""
        teachers = []
        semester = ""
        l1 = []
        global sub_code
        for n, i in enumerate(details):
            flag = 0
            if (len(i) > 1):
                for n1, j in enumerate(i):

                    if (j.startswith("Department")):
                        print("THis is department")
                        print(department)
                        j = j.split("\n")
                        department = j[0]
                        exam = j[1]
                        print(j)
                        print(department)
                    if (i[n1].startswith("Degree")):
                        branch = i[n1 + 1]
                        print(branch)

                    if (i[n1].startswith("Subject")):
                        subject = i[n1 + 1]
                        print(subject)
                    if ("Prepared By" in i[n1]):
                        # print("teachers")
                        if i[n1].endswith("By"):
                            # print("teachers")
                            teachers = details[n + 1][n1]
                        else:
                            teachers = i[n1].split("\n")
                        print(teachers)
                    # print(teachers)
                    if (i[n1].startswith("Date")):
                        date = i[n1 + 1]
                        print(date)

        sheet1.write(0, 4, college)
        sheet1.write(1, 4, "Degree : ")
        sheet1.write(1, 5, branch)
        sheet1.write(2, 4, "Subject : ")
        sheet1.write(2, 5, subject)
        sheet1.write(3, 4, "Date : ")
        sheet1.write(3, 5, date)
        if (type(teachers) == list and len(teachers) > 1):
            sheet1.write(4, 4, "Faculty : ")
            t1 = "".join(teachers)
            sheet1.write(4, 5, t1)
            # for n2, t in enumerate(teachers):

            #   sheet1.write(4, 5 + n2+1, t)
        else:
            sheet1.write(4, 4, "Faculty : ")
            sheet1.write(4, 5, teachers)
        print(subject[:7])
        sub_code = subject[:7]


    wb = Workbook()
    sheet1 = wb.add_sheet('Sheet 1')
    print("SUBJECT CODE")
    print(sub_code)
    sheet1.write(2 + 8, 0, "Q no")
    first_col = sheet1.col(0)
    first_col.width = 256 * 10
    first_col.width
    sheet1.write(2 + 8, 1, "Sub Qn")
    first_col = sheet1.col(1)
    first_col.width = 256 * 5
    first_col.width
    sheet1.write(2 + 8, 2, "CO")
    first_col = sheet1.col(2)
    first_col.width = 256 * 5
    first_col.width
    sheet1.write(2 + 8, 3, "Marks")
    first_col = sheet1.col(3)
    first_col.width = 256 * 5
    first_col.width
    sheet1.write(2 + 8, 4, "Question")
    first_col = sheet1.col(4)
    first_col.width = 256 * 50
    first_col.width

    sheet1.write(2 + 8, 5, "K Level given by faculty")
    sheet1.write(2 + 8, 6, "K Level according to Blooms")
    sheet1.write(2 + 8, 7, "Suggested  action verbs")
    first_col = sheet1.col(7)
    first_col.width = 256 * 70
    first_col.width
    f1 = open("k1.txt", "r")
    f2 = open("k2.txt", "r")
    f3 = open("k3.txt", "r")
    f4 = open("k4.txt", "r")
    f5 = open("k5.txt", "r")
    f6 = open("k6.txt", "r")
    k1_1 = f1.read().split("\n")
    k2_2 = f2.read().split("\n")
    k3_3 = f3.read().split("\n")
    k4_4 = f4.read().split("\n")
    k5_5 = f5.read().split("\n")
    k6_6 = f6.read().split("\n")
    k1 = []
    k2 = []
    k3 = []
    k4 = []
    k5 = []
    k6 = []
    for n, i in enumerate(k1_1):
        if (i == '' or i == ' '):
            continue
        else:
            i = i.lower()
            i = i.strip()
            k1.append(" " + i + " ")

    for n, i in enumerate(k2_2):
        if (i == '' or i == ' '):
            continue
        else:
            i = i.lower()
            i = i.strip()
            k2.append(" " + i + " ")
    for n, i in enumerate(k3_3):
        if (i == '' or i == ' '):
            continue
        else:
            i = i.lower()
            i = i.strip()
            k3.append(" " + i + " ")
    for n, i in enumerate(k4_4):
        if (i == '' or i == ' '):
            continue
        else:
            i = i.lower()
            i = i.strip()
            k4.append(" " + i + " ")
    for n, i in enumerate(k5_5):
        if (i == '' or i == ' '):
            continue
        else:
            i = i.lower()
            i = i.strip()
            k5.append(" " + i + " ")
    for n, i in enumerate(k6_6):
        if (i == '' or i == ' '):
            continue
        else:
            i = i.lower()
            i = i.strip()
            k6.append(" " + i + " ")
    print("K1")
    print(k1)
    print("K2")
    print(k2)
    print("K3")
    print(k3)
    print("K4")
    print(k4)
    print("K5")
    print(k5)
    print("K6")
    print(k6)
    for n, i in enumerate(k1):
        k1[n] = i.strip()


    def maxim(l):
        m = 0
        m1 = "K1"
        for n, j in enumerate(l):
            if (j > m):
                m = j
                m1 = "K" + str(n + 1)
        if (m == 0):
            # return "K1"
            return "Nil"
        return m1


    def mcq(filePath):
        fp = filePath
        document = docx.Document(fp)
        opt = -1
        for n, para in enumerate(document.paragraphs):

            if ("MCQ".casefold() in para.text.casefold() or "Multiple choice questions".casefold() in para.text.casefold()):

                if ("PART" in document.paragraphs[n - 1].text or "Part" in document.paragraphs[n - 1].text):
                    # print(document.paragraphs[n - 1].text)

                    if (" " in document.paragraphs[n - 1].text):
                        l1 = document.paragraphs[n - 1].text.split(" ")
                    elif ("-" in document.paragraphs[n - 1].text):
                        l1 = document.paragraphs[n - 1].text.split("-")
                    # print(l1)
                    for n, i in enumerate(l1):
                        if ('A' == i or 'B' == i or 'C' == i):
                            part = l1[n]
                            # print(part)

                if ("Part".casefold() in para.text.casefold() and (
                        "MCQ".casefold() in para.text.casefold() or "Multiple choice questions".casefold() in para.text.casefold())):

                    check_pt = para.text[para.text.casefold().rindex("part".casefold()):para.text.casefold().rindex(
                        "MCQ".casefold())]

                    for i in ['A', 'B', 'C']:
                        if "- %s" % i in check_pt or " %s " % i in check_pt or "part - %s" % i in check_pt:
                            part = i
                            # print(part)
                if (part == 'A'):
                    opt = 1
                elif (part == 'B'):
                    opt = 2
                elif (part == 'C'):
                    opt = 3
                # print(opt)
                return opt


    def suggest(i):
        a, b, c, d, e, f = i

        if a != f:
            if (a == "K1"):
                k11 = []
                for i in k1:
                    k11.append(i.strip().capitalize())
                return (", ".join(k11))
            if (a == "K2"):
                k22 = []
                for i in k2:
                    k22.append(i.strip().capitalize())
                return (", ".join(k22))
            if (a == "K3"):
                k33 = []
                for i in k3:
                    k33.append(i.strip().capitalize())
                return (", ".join(k33))
            if (a == "K4"):
                k44 = []
                for i in k4:
                    k44.append(i.strip().capitalize())
                return (", ".join(k44))


    def writing_xl(n, i, subq, sub=0):
        nq = sub
        if n in subq:
            mark = i[-1]
            i.pop()

            for n3, q in enumerate(i):
                a, b, c, d, e, f, g, h = q
                sheet1.write(n + n3 + 2 + nq + 8, 0, n)
                sheet1.write(n + n3 + 2 + nq + 8, 1, chr(96 + n3 + 1))
                sub = sub + 1
                sheet1.write(n + n3 + 2 + nq + 8, 2, c)
                sheet1.write(n + n3 + 2 + nq + 8, 3, int(h))
                sheet1.write(n + n3 + 2 + nq + 8, 4, b)
                f = f.upper()
                sheet1.write(n + n3 + 2 + nq + 8, 5, a)
                sheet1.write(n + n3 + 2 + nq + 8, 6, f)
                sheet1.write(n + n3 + 2 + nq + 8, 7, g)

        else:
            a, b, c, d, e, f, g, h = i
            sheet1.write(n + 2 + nq + 8, 0, n)
            sheet1.write(n + 2 + nq + 8, 2, c)
            sheet1.write(n + 2 + nq + 8, 3, int(h))
            sheet1.write(n + 2 + nq + 8, 4, b)
            f = f.upper()
            sheet1.write(n + 2 + nq + 8, 5, a)
            sheet1.write(n + 2 + nq + 8, 6, f)
            sheet1.write(n + 2 + nq + 8, 7, g)
        return sub

    def print_q_mcq(q):
        print("THIS IS SSSSSSSSSS")
        #print("PRINTING MCQS..............")
        l1 = q.casefold().split("(OPTIONS)".casefold())
        print(l1)
        print(l1[0])
        return (l1[0])
    
    def p_q(q):
        print("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@")
        l=[]
        l1 = q.split("\n")
        for i in range (len(l1)):
            if(l1[i]!=""):
                l.append(l1[i])
        print(l)
        #l=len(l)
        
        if("Answer" in l[len(l)-1]):
            s=" ".join(l[:-5])
        else:
            s=" ".join(l[:-4])
            
        print("THIS IS SSSSSSSSSS")
        print(s)
        
        return(s)
    

    def print_q(opt):
        n = 1
        i = opt - 1
        for j in range(len(tables[i])):
            l1 = (tables[i][n][1].split("\n"))
            return l1[0]
            n += 1


    def classify_k(K_Lev, question):
        print(K_Lev, question)
        question = question.lower()
        l = []
        l2 = []
        q = question
        c_k1 = []
        c_k2 = []
        c_k3 = []
        c_k4 = []
        punct = [',', ':', '.', '!', ';', '\\', '-', '$', '(', ')']

        q = q.replace('\n', ' ')

        for i in question:
            if i in punct:
                q = q.replace(i, ' ')
        q = " " + q

        print(q)
        for n, i in enumerate(k1):
            if search(k1[n], q):

                if K_Lev == 'K1':
                    c_k1.append("Chosen 500")
                elif (i not in c_k1):
                    c_k1.append(i)

        for n, i in enumerate(k2):
            if search(k2[n], q):
                if K_Lev == 'K2':
                    c_k2.append("Chosen 500")
                elif (i not in c_k2):
                    c_k2.append(i)

        for n, i in enumerate(k3):
            if search(k3[n], q):
                if K_Lev == 'K3':
                    c_k3.append("Chosen 500")
                elif (i not in c_k3):
                    c_k3.append(i)
        for n, i in enumerate(k4):
            if search(k4[n], q):
                if K_Lev == 'K4':
                    c_k4.append("Chosen 500")
                elif (i not in c_k4):
                    c_k4.append(i)
        l.append(c_k1)
        l.append(c_k2)
        l.append(c_k3)
        l.append(c_k4)

        l2.append(len(c_k1))
        l2.append(len(c_k2))
        l2.append(len(c_k3))
        l2.append(len(c_k4))

        for i in range(len(l)):
            if l[i] == ["Chosen 500"]:
                l[i] = []
                l2[i] = 500
        #print("l and l2 HERE ", l, l2)
        return l, l2


    row_data = {}
    data = []
    kfinder = lambda x: "".join([let for let in x if let == 'K' or let in ['1', '2', '3', '4', '5',
                                                                           '6']])  # lambda function to find the k values in a pattern like <K2> etc
    cofinder = lambda x: "".join([let for let in x if let in 'CO' or let in ['1', '2', '3', '4', '5',
                                                                             '6']])

    def sub_marks(word, blis=[]):
        maxmark = 10
        msub = []
        #print("OVER HERE",word)
        
        res = [j for j in range(len(word)) if word.startswith('(', j)]
        for j in res:
            k = word[j:].index(')') + len(word[:j])
            q = j + 1
            while word[q] == ' ':
                q += 1
            v = j + 1
            if q != v and q != k:
                j = q - 1
                v = q

            if v == k:  # TO HANDLE () AND TO MAKE SURE THAT IT IS
                v = k + 2

            while word[v].isdigit() == True or word[v] == '+' and v < k:
                v += 1

            if (v == k or 'Marks'.casefold() in word[v:k].casefold()):
                if '+' in word[j + 1:v]:
                    if int(word[j + 1]) + int(word[j + 3]) < maxmark:
                        msub.append(int(word[j + 1]) + int(word[j + 3]))
                        blis.append(k)
                elif int(word[j + 1:v]) < maxmark and int(word[j + 1:v]) != 0:
                    msub.append(int(word[j + 1:v]))
                    blis.append(k)
            
        return msub


    def subquestions(subq):  # Returns the subquestions as as a list of lists
        msub = []
        blis = []
        temp = []
    
        for typ, i in enumerate(subq):
            if typ == 1:

                l = len(sub_marks(subq[1], blis))
                te = []
                spn = 0
                for br in range(len(blis)):
                    if br == 0:
                        te.append(subq[1][:blis[br] + 1].replace('\n', ' '))
                        spn = blis[br] + 1
                    elif br == len(blis) - 1:
                        if blis[br] != len(subq[1]):
                            te.append(subq[1][spn:].replace('\n', ' '))
                        else:
                            te.append(subq[1][spn:blis[br] + 1].replace('\n', ' '))

                    else:
                        te.append(subq[1][spn:blis[br] + 1].replace('\n', ' '))
                        spn = blis[br] + 1

                i = te
            else:
                i = i.split("\n")

            i = [i[k] for k in range(len(i)) if i[k] != '']  # To remove the spaces

            temp.append(i)
        temp[0] = list(map(kfinder, temp[0]))
        temp[2] = list(map(cofinder, temp[2]))
        for i in range(len(temp[2])):  # ADD THIS TO CODE
            if temp[2][i] == '':
                temp[2].remove('')  # ADD THIS TO CODE
        if len(temp[0]) != l:
            temp[0] = temp[0] * l
        if len(temp[2]) != l:
            temp[2] = temp[2] * l

        subq = [list(each) for each in list(zip(temp[0], temp[1], temp[2]))]

        return subq


    tables, oqn, sub_q = table_text(filePath)
    fp = filePath
    for l in range(len(tables)):
        for qno, i in tables[l].items():
            if qno not in sub_q.keys():
                '''
                if (mcq(fp) == -1 or mcq(fp) == None):
                    j, s = classify_k(i[0], i[1])
                    i.append(j)
                    i.append(s)
                    i.append(maxim(i[4]))
                    w = suggest(i)
                    i.append(w)
                '''    

                if(mcq(fp)==l+1):
                    opt = mcq(fp)
                    mcq_question = print_q_mcq(i[1])
                    print("###############################################")
                    print(mcq_question)
                    j, s = classify_k(i[0], mcq_question)
                    i.append(j)
                    i.append(s)
                    i.append(maxim(i[4]))
                    w = suggest(i)
                    i.append(w)
                else:
                    j, s = classify_k(i[0], i[1])
                    i.append(j)
                    i.append(s)
                    i.append(maxim(i[4]))
                    w = suggest(i)
                    i.append(w)


            else:
                for i in tables[l][qno]:
                    j, s = classify_k(i[0], i[1])
                    i.append(j)
                    i.append(s)
                    i.append(maxim(i[4]))
                    w = suggest(i)
                    i.append(w)


    def write_tables(tables, subq):
        n = 1
        na = 0
        for i in range(len(tables)):  # 3 for CAT 1 doc
            for j in range(len(tables[i])):
                na = writing_xl(n, tables[i][n], subq, na)
                n += 1


    def print_tab(tables):  # REMOVE in final
        # print("TABLE IS BEING PRINTED@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@")
        n = 1
        for i in range(len(tables)):  # 3 for CAT 1 doc
            print("Part",i+1)
            for j in range(len(tables[i])):
                print(n, tables[i][n])
                n += 1


    def subq_marks(KL, CO, row, qno, mark_scheme, subq):
        # mark_scheme = {0: 2, 1: 10, 2: 10}
        marno = 0
        for each in row[qno]:
            if isinstance(each, int) == False:
                # if each[5] != "Nil":
                KL[each[5]] = KL[each[5]] + subq[qno][marno]
                CO[each[2]] = CO[each[2]] + subq[qno][marno]
                marno += 1
        return KL, CO


    def mark_app(tables, subq, d={'6': '2', '3': '6', '2': '10'}):
        qno = 1
        mark = 0
        for i in range(len(tables)):
            for n1, k in enumerate(d.keys()):
                if (n1 == i):
                    # print(n1)
                    mark = d[k]

            for j in range(len(tables[i])):
                if qno not in subq.keys():
                    tables[i][qno].append(int(mark))

                else:
                    for a in range(len(tables[i][qno])):
                        tables[i][qno][a].append(int(subq[qno][a]))
                    tables[i][qno].append(int(mark))

                qno += 1


    def format_one_of(temp):
        i = 0
        st = ""
        while i < len(temp):  # ADDDDDD
            if temp[i] == 'K1' or temp[i] == 'K2':
                st = st + " K1/K2"
            else:
                st = st + " K3/K4"
            if i == len(temp) - 1:
                st = st + " ]"
            else:
                st = st + " ,"
            i += 1
        return st


    def finding_possib(levels, num):  
        possib = []
        if levels == 0:
            return 0
        for i in range(len(levels) + 1):
            for j in range(i):
                if (len(levels[j:i]) == num):
                    possib.append(levels[j:i])
        possib = list(set(map(tuple, possib)))
        return possib


    def maxpart(ele, lowerPartA, upperPartA, tot, mar):  
        max_k3 = 0
        max_k1 = 0

        for i in ele:
            if i == 'K2' or i == 'K1':
                max_k1 += mar  # add the respective mark for that section
            else:
                max_k3 += mar

        max_k3 += upperPartA
        max_k1 += lowerPartA
        max_k1 = max_k1 / tot * 100
        max_k3 = max_k3 / tot * 100
        return max_k1, max_k3


    def letsee(s3of5, subq, choice):
        print("-----------------------------------")
        print("WITHIN LETSSEEEEE")

        temp = []
        global AddComments, lower_max_k1, upper_max_k3
        parts = {0: 'A', 1: 'B', 2: 'C'}
        levels = []
        s = ""
        # choice=[1, 2, 3, [10, 20], 50, 10]
        # levels = ['K3', 'K1', 'K1']

        err = 0
        dpa, spa = calc_am((choice[3][0] / choice[4]) * 100, -1, 0)
        s = s + "Checking the Choose any %d out of %d in Part %s." % (choice[1], choice[2], parts[choice[0]])
        if "max" in spa and dpa == False:
            print("ERROR 1")
            print(
                "The total marks for K1/K2 in all the parts except for Part B exceed the maximum marks allowed. In the paper you can have a maximum of %d marks for K1 and K2. We have %d marks. The marks for K1 and K2 must be reduced by %d" % (
                    (lower_max_k1 / 100) * choice[4], choice[3][0], choice[3][0] - ((lower_max_k1 / 100) * choice[4])))
            err += 1
            s = s + "The total marks for K1/K2 in all the parts except for Part B exceed the maximum marks allowed.In the paper you can have a maximum of %d marks for K1 and K2.We have %d marks.The marks for K1 and K2 must be reduced by %d" % (
                (lower_max_k1 / 100) * choice[4], choice[3][0], choice[3][0] - ((lower_max_k1 / 100) * choice[4]))


        dpa, spa = calc_am(-1, (choice[3][1] / choice[4]) * 100, 0)
        if "max" in spa and dpa == False:
            print("ERROR 3")
            print("The total marks for K3/K4 in all the parts except for Part B exceed the maximum marks allowed. "
                  "In the paper you can have a maximum of %d marks for K3 and K4. "
                  "We have %d marks. The marks for K3 and K4 must be reduced by %d" % (
                      (upper_max_k3 / 100) * choice[4], choice[3][1],
                      choice[3][1] - ((upper_max_k3 / 100) * choice[4])))
            s = s + "The total marks for K3/K4 in all the parts except for Part B exceed the maximum marks allowed.In the paper you can have a maximum of %d marks for K3 and K4.We have %d marks.The marks for K3 and K4 must be reduced by %d" % (
                (upper_max_k3 / 100) * choice[4], choice[3][1], choice[3][1] - ((upper_max_k3 / 100) * choice[4]))
            err += 1

        if err == 0:
            for qno, each in s3of5.items():
                print(each)
                if qno not in subq.keys():
                    if isinstance(each, int) == False and each[5] != 'Nil':
                        levels.append(each[5])
                else:
                    for ev in each:
                        if isinstance(ev, int) == False and ev[5] != 'Nil':
                            temp.append(ev[5])
                    temp = list(set(temp))
                    if len(temp) == 1:
                        levels.append(temp[0])
                    elif len(temp) != 0:
                        levels.append(max(temp))

            #print("1 Currently before everything levels is ", levels)

            temp = sorted(levels)
            possib = finding_possib(temp, choice[1])
            for i in range(len(temp)):
                if temp[i] == 'K2':
                    temp[i] = 'K1'
                elif temp[i] == 'K4':
                    temp[i] = 'K3'

            print("TEMP RN IS BEFORE LOOP ", temp)
            print('possib BEFORE LOOP', possib)

            flagb = 0
            count = 1

            cklev = {'K1': levels.count('K1') + levels.count('K2'), 'K3': levels.count('K3') + levels.count('K4')}
            max_k1, max_k3 = maxpart(possib[0], choice[3][0], choice[3][1], choice[4], choice[5])
            #print('ckkk', cklev)
            #print("MAAA", max_k1, max_k3)
            checkforprob = []
            checkforprob.append(temp)

            while (flagb != -1 and (checkforprob.count(
                    temp) != 2 or count == 1) and len(temp) == choice[
                       2]):

                for ele in possib:
                    max_k1, max_k3 = maxpart(ele, choice[3][0], choice[3][1], choice[4], choice[5])
                    if calc_am(max_k1, max_k3, 1) == "no":
                        flagb += 1
                        dist1, s1 = calc_am(max_k1, -1, 0)
                        dist2, s2 = calc_am(-1, max_k3, 0)
                        if dist1 == False:
                            if "max" in s1:
                                cklev['K1'] -= 1
                                cklev['K3'] += 1
                            elif "min" in s1:
                                cklev['K1'] += 1
                                cklev['K3'] -= 1
                        if dist2 == False and dist1 == True:
                            if "max" in s2:
                                cklev['K3'] -= 1
                                cklev['K1'] += 1
                            elif "min" in s2:
                                cklev['K3'] += 1
                                cklev['K1'] -= 1
                    #else:
                        #print("for this ele ", ele, "it works")
                    #print('In looppp ckkk', cklev)

                count += 1
                if flagb == 0:
                    flagb = -1
                else:
                    flagb = 0
                temp = ['K1'] * cklev['K1'] + ['K3'] * cklev['K3']

                possib = finding_possib(temp, choice[1])
                checkforprob.append(temp)
                print("CHECKS FDO", checkforprob)
                print("T", temp)
                #print("possy", possib)

            if flagb == -1 and temp == checkforprob[0]:
                print("The given condn is satisfied. So this list is fine")
                s = s + "No changes have to be made."
            elif flagb == -1:
                print("The given list is wrong. The right list is ", temp)
                s = s + ".You cannot have [ "
                for i in range(len(levels)):
                    if i != len(levels) - 1:
                        s = s + levels[i] + " , "
                    else:
                        s = s + levels[i]
                s = s + " ]." + "Instead you can have [%s" % format_one_of(temp)
            elif (checkforprob.count(temp) >= 2) or (len(temp) != choice[2]):  # temp in checkforprob:
                print("There is something wrong with the percentage guideline")
                s = s + "There is something wrong with the percentage guidelines.Regardless of the K levels chosen for this Part the guidelines will not be satisfied.Please change the percentage guidelines."
        print("\n\n~~~~~~ At the end of this letsee function this is the string ~~~~~~")
        print(s,"\n")
        AddComments.append([s])

    def content(fp):  # to obtain the details from the user
        document = docx.Document(fp)
        choice = [-1, -1]
        l1 = []
        l2 = []
        part = ''
        numb1 = {("one", "1"): 1, ("two", "2"): 2, ("three", "3"): 3, ("four", "4"): 4}
        for n, para in enumerate(document.paragraphs):

            if ("Answer any".casefold() in para.text.casefold() or "any".casefold() in para.text.casefold()) and (
                    "questions".casefold() in para.text.casefold() or "out of".casefold() in para.text.casefold()):

                if ("PART" in document.paragraphs[n - 1].text or "Part" in document.paragraphs[n - 1].text):
                    print(document.paragraphs[n - 1].text)

                    if (" " in document.paragraphs[n - 1].text):
                        l1 = document.paragraphs[n - 1].text.split(" ")
                    elif ("-" in document.paragraphs[n - 1].text):
                        l1 = document.paragraphs[n - 1].text.split("-")
                    print(l1)
                    for n, i in enumerate(l1):
                        if ('A' == i or 'B' == i or 'C' == i):
                            part = l1[n]
                            choice[0] = l1[n]
                            print(part)
                            # return 8

                if ("Part".casefold() in para.text.casefold() and "Any".casefold() in para.text.casefold()):

                    check_pt = para.text[para.text.casefold().rindex("part".casefold()):para.text.casefold().rindex(
                        "any".casefold())]

                    for i in ['A', 'B', 'C']:
                        if "- %s" % i in check_pt or " %s " % i in check_pt or "part - %s" % i in check_pt:
                            choice[0] = i

                check_pt = para.text[para.text.casefold().rindex("any".casefold()):para.text.casefold().rindex(
                    "(".casefold())].casefold()

                for i in numb1.keys():
                    for each in i:
                        if each.casefold() in check_pt:
                            choice[1] = each
                k = 'A'
                while k <= 'C':
                    if choice[0] == k:
                        choice[0] = ord(k) - 65
                    k = chr(ord(k) + 1)

                for key, val in numb1.items():
                    if choice[1] in key:
                        choice[1] = val
        if (choice.count(-1) == 1):
            choice = [-1, -1]

        return choice


    def calc_am(dist1=-1, dist2=-1, flag=1):  # 

        global lower_max_k1
        global lower_min_k1
        global upper_min_k3
        global upper_max_k3
        temp = []
        s = ""

        if flag == 1 or flag == 0:
            if lower_min_k1 != 0 and dist1 != -1:
                if dist1 >= lower_min_k1:
                    s = s + "min k1-"
                    temp.append(True)
                else:  # check <= or <
                    s = s + "min k1-"
                    temp.append(False)
            elif lower_max_k1 != 0 and dist1 != -1:
                if dist1 <= lower_max_k1:
                    s = s + "max k1-"
                    temp.append(True)
                else:
                    s = s + "max k1-"
                    temp.append(False)
            if upper_max_k3 != 0 and dist2 != -1:
                if dist2 <= upper_max_k3:
                    s = s + "max k3"
                    temp.append(True)
                else:
                    s = s + "max k3"
                    temp.append(False)
            elif upper_min_k3 != 0 and dist2 != -1:
                if dist2 >= upper_min_k3:
                    s = s + "min k1"
                    temp.append(True)
                else:
                    s = s + "min k1"
                    temp.append(False)

        elif flag == -1:

            if lower_min_k1 != 0 and dist1 != -1:
                if dist1 < lower_min_k1:
                    temp.append(True)
                else:  # check <= or <
                    temp.append(False)
            elif lower_max_k1 != 0 and dist1 != -1:
                if dist1 > lower_max_k1:
                    temp.append(True)
                else:
                    temp.append(False)
            if upper_max_k3 != 0 and dist2 != -1:
                if dist2 > upper_max_k3:
                    temp.append(True)
                else:
                    temp.append(False)
            elif upper_min_k3 != 0 and dist2 != -1:
                if dist2 < upper_min_k3:
                    temp.append(True)
                else:
                    temp.append(False)
        if dist1 == -1 and dist2 == -1:
            return None
        elif dist1 == -1:
            if flag == 0:
                #print(temp)
                return temp[0], s
            return temp[0]
        elif dist2 == -1:
            if flag == 0:
                return temp[0], s
            return temp[0]

        else:
            if temp[0] and temp[1] == True:
                return "yes"
            else:
                return "no"


    def markdist(tables, subq, oqn, fp, nooq, mark_scheme):  # FINAL CHANGE

        init_lower = 0
        init_upper = 0
        CO = {}
        KL = {"Nil": 0}
        orq = []
        global AddComments
        for j in oqn:
            for eq in range(len(j)):
                if eq != 0:
                    orq.append(j[eq])

        choices = content(fp)
        #print("CHOICES ARE", choices)

        for i in range(-1, len(tables)):  # range(len(tables)):
            for k in range(1, 7):  # Knowledge levels are from 1 to 6
                if i == -1:
                    KL["K%i" % (k)] = 0
                    CO["CO%i" % (k)] = 0
                elif i != choices[0]:

                    KL["K%i" % (k)] = ["Here" for qno, eachl in tables[i].items() if
                                       qno not in subq.keys() and qno not in orq if
                                       "K%i" % (k) in eachl[5]].count("Here") * mark_scheme[i] + KL["K%i" % (k)]
                    CO["CO%i" % (k)] = ["Here" for qno, eachl in tables[i].items() if
                                        qno not in subq.keys() and qno not in orq if
                                        "CO%i" % (k) in eachl[2]].count("Here") * mark_scheme[i] + CO["CO%i" % (k)]
                if choices != [-1, -1] and i != choices[0] and i != -1:  
                    #print("Entered", choices[0], i)
                    if k == 1 or k == 2:
                        init_lower += ["Here" for qno, eachl in tables[i].items() if
                                       qno not in subq.keys() and qno not in orq if
                                       "K%i" % (k) in eachl[5]].count("Here") * mark_scheme[i]
                    elif k == 3 or k == 4:
                        init_upper += ["Here" for qno, eachl in tables[i].items() if
                                       qno not in subq.keys() and qno not in orq if
                                       "K%i" % (k) in eachl[5]].count("Here") * mark_scheme[i]

                    #print("init low", init_lower, "up", init_upper)
            if i != -1 and i != choices[0]:
                KL["Nil"] = ["Here" for qno, eachl in tables[i].items() if qno not in subq.keys() and qno not in orq if
                             eachl[5] == "Nil"].count("Here") * mark_scheme[i] + KL["Nil"]

            #print(KL)
            if i == 0 and len(tables[i]) != nooq[i]:
                divd = nooq[i] / len(tables[i])
                for key, val in KL.items():
                    KL[key] = int(val * divd)
                for key, val in CO.items():
                    CO[key] = int(val * divd)
                if choices != [-1, -1] and i < choices[0] and i != -1:  
                    #print("Entered diss", choices[0], i)
                    init_lower = int(init_lower * divd)
                    init_upper = int(init_upper * divd)
                    #print("init low", init_lower, "up", init_upper)
                    

            #print(KL)
            #print("KL123456", init_lower, init_upper)

        for eq in oqn:
            if eq[1] in subq.keys() or eq[2] in subq.keys():  # Checking if any of the or questions have subquestions

                q1 = [q[5] for q in tables[eq[0] - 1][eq[1]] if isinstance(q,
                                                                           int) == False]  # and q[5] != "Nil"]   #Here Nil is necessary because you're checking distribution of k levels. But it will still add the co if there is a nil
                #print("WE ARE PRINTING EQ ,Q1,TABLES ",eq,q1,tables[eq[0]-1][eq[2]])
                q2 = [q[5] for q in tables[eq[0] - 1][eq[2]] if
                      isinstance(q, int) == False]  # and q[5] != "Nil"]  
                for eve in range(len(eq[1:])):
                    if eve == 0:
                        dist1 = (q1.count('K1') + q1.count('K2')) / len(q1) * 100
                        dist2 = (q1.count('K3') + q1.count('K4')) / len(q1) * 100
                        q1.append(calc_am(dist1, dist2, 1))

                    else:
                        dist1 = (q2.count('K1') + q2.count('K2')) / len(q2) * 100
                        dist2 = (q2.count('K3') + q2.count('K4')) / len(q2) * 100

                        q2.append(calc_am(dist1, dist2, 1))

                if q1[len(q1) - 1] == 'Yes' and q2[len(q2) - 1] == 'No':

                    KL, CO = subq_marks(KL, CO, tables[eq[0] - 1], eq[1], mark_scheme, subq)

                elif q1[len(q1) - 1] == 'No' and q2[len(q2) - 1] == 'Yes':

                    KL, CO = subq_marks(KL, CO, tables[eq[0] - 1], eq[2], mark_scheme, subq)

                else:

                    KL, CO = subq_marks(KL, CO, tables[eq[0] - 1], eq[1], mark_scheme, subq)

            else:
                q1 = tables[eq[0] - 1][eq[1]][5]
                q2 = tables[eq[0] - 1][eq[2]][5]
                c1 = tables[eq[0] - 1][eq[1]][2]
                c2 = tables[eq[0] - 1][eq[2]][2]
                if q1 == q2 or (q1 in ['K1', 'K2'] and q2 in ['K1', 'K2']) or (
                        q1 in ['K3', 'K4'] and q2 in ['K3', 'K4']) or calc_am(mark_percentage(KL)["12"],
                                                                              mark_percentage(KL)["34"], 1) == "yes":
                    if q1 != 'Nil':  
                        KL[q1] = KL[q1] + mark_scheme[eq[0] - 1]
                        CO[c1] = CO[c1] + mark_scheme[eq[0] - 1]
                    elif q2 != 'Nil':
                        KL[q2] = KL[q2] + mark_scheme[eq[0] - 1]
                        CO[c2] = CO[c2] + mark_scheme[eq[0] - 1]
                    else:
                        KL[q1] = KL[q1] + mark_scheme[eq[0] - 1]
                        CO[c1] = CO[c1] + mark_scheme[eq[0] - 1]  
                else:
                    if (calc_am(mark_percentage(KL)["12"], -1, -1)):
                        if q1 in ['K3', 'K4']:
                            print("1 markdist")
                            KL[q1] = KL[q1] + mark_scheme[eq[0] - 1]
                            CO[c1] = CO[c1] + mark_scheme[eq[0] - 1]  

                        else:
                            if q2 != 'Nil':
                                print("2 markdist")
                                KL[q2] = KL[q2] + mark_scheme[eq[0] - 1]
                                CO[c2] = CO[c2] + mark_scheme[eq[0] - 1]  
                            else:
                                print("3 markdist")
                                KL[q1] = KL[q1] + mark_scheme[eq[0] - 1]
                                CO[c1] = CO[c1] + mark_scheme[eq[0] - 1]


                    elif (calc_am(-1, mark_percentage(KL)["34"], -1)):
                        if q1 in ['K1', 'K2']:
                            print("4 markdist")
                            KL[q1] = KL[q1] + mark_scheme[eq[0] - 1]
                            CO[c1] = CO[c1] + mark_scheme[eq[0] - 1]  
                        else:

                            if q2 != 'Nil':
                                print("5 markdist")
                                KL[q2] = KL[q2] + mark_scheme[eq[0] - 1]
                                CO[c2] = CO[c2] + mark_scheme[eq[0] - 1]
                            else:
                                print("6 markdist")
                                KL[q1] = KL[q1] + mark_scheme[eq[0] - 1]
                                CO[c1] = CO[c1] + mark_scheme[eq[0] - 1]

        for i in range(len(tables)):  # TO HANDLE THE SUBQUESTIONS WHICH ARE NOT AN OR QUESTION
            for qno in subq:
                #print("KL", KL, init_lower, init_upper)
                if qno not in orq:
                    marno = 0

                    if qno in tables[i].keys():

                        for each in tables[i][qno]:
                            if isinstance(each, int) == False and i != choices[0]:
                                #print(each[5])
                                KL[each[5]] = KL[each[5]] + subq[qno][marno]
                                CO[each[2]] = CO[each[2]] + subq[qno][marno]
                                if choices != [-1, -1]:
                                    if each[5] in ['K1', 'K2'] and i != choices[
                                        0]:  # to handle the choose any _ out of _ questions
                                        init_lower += subq[qno][marno]
                                    elif each[5] in ['K3', 'K4'] and i != choices[0]:
                                        init_upper += subq[qno][marno]
                                marno += 1

        if choices != [-1, -1]:  # Add the nil condition
            choices.append(len(tables[choices[0]]))
            choices.append([init_lower, init_upper])
            qno = 0
            for i in range(choices[0]):
                qno = qno + len(tables[i])
            qno += 1
            for i in range(choices[1]):
                if qno in subq.keys():
                    subq_marks(KL, CO, tables[choices[0]], qno, mark_scheme, subq)
                else:
                    KL[tables[choices[0]][qno][5]] = KL[tables[choices[0]][qno][5]] + mark_scheme[choices[0]]
                    CO[tables[choices[0]][qno][2]] = CO[tables[choices[0]][qno][2]] + mark_scheme[choices[0]]
                qno += 1

            choices.append(sum(KL.values()))
            choices.append(mark_scheme[choices[0]])
            print("choices rn are", choices)
            if KL["Nil"] == 0:
                letsee(tables[choices[0]], subq, choices)
            else:
                AddComments.append(["To see if the KL levels for (Choose any %d out of %d questions) type questions follow the given guidelines.\nPlease add the necessary action verbs or change the K levels for all the questions " % (
                        choices[1], choices[2])])

        return KL, CO


    def mark_percentage(KL):  # maybe have to make indices more general incase you plan to use it for CO also
        dist = {}
        dist["12"] = (KL['K1'] + KL['K2']) / sum(KL.values()) * 100
        dist["34"] = (KL['K3'] + KL['K4']) / sum(KL.values()) * 100
        return dist

    def Course_check(tables, orqn, subq):  # Checking whether the course level for OR questions is the same

        for i in range(len(orqn)):
            if orqn[i][1] not in subq.keys() and orqn[i][
                2] not in subq.keys():  # Checking whether the or questions are also have subquestions
                if tables[orqn[i][0] - 1][orqn[i][1]][2] == tables[orqn[i][0] - 1][orqn[i][2]][2]:
                    orqn[i].append(
                        ['CO', True, tables[orqn[i][0] - 1][orqn[i][1]][2], tables[orqn[i][0] - 1][orqn[i][2]][2]])
                else:

                    orqn[i].append(
                        ['CO', False, tables[orqn[i][0] - 1][orqn[i][1]][2], tables[orqn[i][0] - 1][orqn[i][2]][2]])
                if tables[orqn[i][0] - 1][orqn[i][1]][5] == tables[orqn[i][0] - 1][orqn[i][2]][
                    5]:  # HAVE TO CHECK FOR NIL HERE FOR THE K LEVELLS
                    orqn[i].append(
                        ['KL', True, tables[orqn[i][0] - 1][orqn[i][1]][5], tables[orqn[i][0] - 1][orqn[i][2]][5]])

                else:
                    orqn[i].append(
                        ['KL', False, tables[orqn[i][0] - 1][orqn[i][1]][5], tables[orqn[i][0] - 1][orqn[i][2]][5]])

            else:
                if tables[orqn[i][0] - 1][orqn[i][1]][0][2] == tables[orqn[i][0] - 1][orqn[i][2]][0][2]:
                    orqn[i].append(
                        ['CO', True, tables[orqn[i][0] - 1][orqn[i][1]][0][2],
                         tables[orqn[i][0] - 1][orqn[i][2]][0][2]])
                else:
                    orqn[i].append(
                        ['CO', False, tables[orqn[i][0] - 1][orqn[i][1]][0][2],
                         tables[orqn[i][0] - 1][orqn[i][2]][0][2]])
                q1 = [q[5] for q in tables[orqn[i][0] - 1][orqn[i][1]] if
                      isinstance(q, int) == False and q[5] != "Nil"]  
                q2 = [q[5] for q in tables[orqn[i][0] - 1][orqn[i][2]] if
                      isinstance(q, int) == False and q[5] != "Nil"]  
                #print(q1, q2)
                if max(q1) == max(q2):  # and min(q1) == min(q2):
                    orqn[i].append(['KL', True, max(q1), max(q2)])
                else:
                    orqn[i].append(['KL', False, max(q1), max(q2)])

        return orqn


    def makemark(mnd):
        d = {}
        noq = {}
        no = 0
        tot_marks = 0
        if '' in mnd.keys():
            mnd.pop('')
        for key, val in mnd.items():
            if key != '':
                d[no] = int(val)
                noq[no] = int(key)
                no += 1
                tot_marks = tot_marks + int(key) * int(val)

        return d, noq, tot_marks


    def add_c(dist1, dist2, orq, KL, fp):
        global AddComments
        global lower_min_k1, lower_max_k1, upper_max_k3, upper_max_k3
        parts = ['N/A', 'A', 'B', 'C']
        st = ""
        print(orq)

        tot = sum(KL.values())
        # st=st+"NIL-There are no action verbs from the Bloom's taxonomy tables present in the question\n "
        if KL["Nil"] != 0:
            st = st + "\nThe questions that do not have any action verbs from Bloom's taxonomy contributes %d marks" % \
                 KL[
                     "Nil"]
            AddComments.append([st])
        else:
            temp, t = calc_am(dist1, -1, 0)

            st = ""
            perc = "Guidelines : "
            temp1, t1 = calc_am(-1, dist2, 0)
            if temp == True and temp1 == True:
                st = st + "Questions Approved.They follow the given guidelines\n"
            else:

                # if temp==True:
                # st=st+"K1 K2 Questions Approved\nIt follows the given guideline of (K1,K2) questions.\n\n"#%(dist1)
                if "max" in t and temp == False:
                    perc = perc + "Max %d %% K1 K2 questions and " % (lower_max_k1)
                    st = st + "\nThe number of K1,K2 questions has to be decreased by %d %% for %d marks. \n\n" % (
                        abs(dist1 - lower_max_k1), abs(dist1 - lower_max_k1) / 100 * tot)
                elif "min" in t and temp == False:
                    perc = perc + "Min %d %% K1 K2 questions and " % (lower_min_k1)
                    st = st + "\nThe number of K1,K2 questions has to be increased by %d %% for %d marks.\n\n" % (
                        abs(dist1 - lower_min_k1), abs(dist1 - lower_min_k1) / 100 * tot)

                # if temp1==True:
                # st=st+"\nK3 K4 Questions Approved.\nNo changes have to be made.\nIt follows the guideline distribution percentages of (K3,K4) questions.\n\n"%(dist2)
                if "max" in t1 and temp1 == False:
                    perc = perc + " Max %d %% K3 K4 questions " % (upper_max_k3)
                    st = st + "\nThe number of K3,K4 questions has to be decreased by %d %% for %d marks.\n\n" % (
                        abs(dist2 - upper_max_k3), abs(dist2 - upper_max_k3) / 100 * tot)
                elif "min" in t1 and temp1 == False:
                    perc = perc + " Min %d %% K3 K4 questions " % (upper_min_k3)
                    st = st + "\nThe number of K3,K4 questions has to be increased by %d %% for %d marks.\n\n" % (
                        abs(dist2 - upper_min_k3), abs(dist2 - upper_min_k3) / 100 * tot)
                AddComments.append([perc])
            AddComments.append([st])
        st = ""
        for eachq in orq:
            d = parts[eachq[0]]
            # if eachq[3][1] == True:
            #    st = st + "\nIn the OR questions %d and %d the CO levels match : NO CHANGES.\n " % (eachq[1], eachq[2])
            if eachq[3][1] == False:
                st = st + "\n In the OR questions %d and %d the CO levels do not match.\n It is %s in %d and %s in %d. \nOr questions should have the same CO level\n" % (
                    eachq[1], eachq[2], eachq[3][2], eachq[1], eachq[3][3], eachq[2])

            # if eachq[4][1] == True:
            #    st = st + "\nIn the OR questions %d and %d the Knowledge levels match: NO CHANGES.\n" % (eachq[1], eachq[2])
            if eachq[4][1] == False:
                st = st + "\n In the OR questions %d and %d the Knowledge levels do not match.\n It is %s in %d and %s in %d. \nOr questions should have the same maximum Knowledge level\n\n" % (
                    eachq[1], eachq[2], eachq[4][2], eachq[1], eachq[4][3], eachq[2])

        choices = content(fp)
        AddComments.append([st])
        if choices != [-1, -1]:
            AddComments.append(AddComments[0])
            AddComments.pop(0)
        if ([''] in AddComments):
            AddComments.remove([''])


    print_tab(tables)
    print("^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^Tables are^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^^")
    print(tables)
    print("Details are")  
    print(details)  
    for i in details:
        print(i)

    #print("Mark distribution is: ", markNumberDict)
    markno, nooq, tot_marks = makemark(markNumberDict)
    #print("The string markno is: ", markno,"\n2. Mark distribution is: ", markNumberDict)
    #print("\nThe total marks are", tot_marks, "\n\n")

    mark_app(tables, sub_q, markNumberDict)  # prints tables with question number
    write_tables(tables, sub_q)

    dept = "CSE"  # changee
    dept1 = dept + ".xls"
    wb.save(dept1)
    print("Sub q is ",sub_q)

    Know_lev, Cour_Out = markdist(tables, sub_q, oqn, filePath, nooq, markno)

    print("Knowledge level distribution", Know_lev)  # prints distribution of knowledge level as dictionary
    print("Course outcome level distribution", Cour_Out)  # prints distribution of course outcome as dictionary
    
    dist = mark_percentage(Know_lev)
    print("K LEVEL:1 and 2 % is", dist["12"])  # prints distribution of knowledge level
    print("K LEVEL:3 and 4 % is", dist["34"])  # prints distribution of course level
    co_ch = Course_check(tables, oqn, sub_q)
    print("The or question CO check is", co_ch)
    add_c(dist["12"], dist["34"], co_ch, Know_lev, filePath)
    print("Sub q is ",sub_q)
    print("AddComments is \n",AddComments)
    
    import openpyxl as xl

    from openpyxl.chart import *


    def openFile(filename):
        os.startfile(filename)

    import xlrd

    from openpyxl.workbook import Workbook
    from openpyxl.reader.excel import load_workbook, InvalidFileException
    import win32com.client as win32


    def save_xl(dept1):
        fp = os.getcwd()
        fname = fp + "/" + dept1
        excel = win32.gencache.EnsureDispatch('Excel.Application')
        wb = excel.Workbooks.Open(fname)
        st1 = str("/")
        st2 = str("\\")
        fname1 = fname + "x"
        fp2 = fname1.replace(st1, st2)
        wb.SaveAs(fp2, FileFormat=51)
        wb.Close()
        excel.Application.Quit()
        return str(dept1 + "x")


    wb = xl.Workbook()
    sheet = wb.active
    wb.save("barChart.xlsx")
    dept2 = save_xl(dept1)

    import openpyxl
    from openpyxl import load_workbook
    from openpyxl.styles import *
    import csv
    import os
    import copy
    from openpyxl.styles.borders import Border, Side

    import xlwings as xw
    from openpyxl.styles import PatternFill, colors
    from openpyxl.styles.differential import DifferentialStyle
    from openpyxl.formatting.rule import Rule
    from openpyxl.formatting.rule import ColorScaleRule

    dept2 = "CSE.xlsx"

    rel_path = os.getcwd()
    fp = rel_path + "\\" + dept2


    def end_coord(workbooklocation=fp, sheetname="Sheet 1", columnletter="A"):
        wb = xw.Book(workbooklocation)
        X = wb.sheets[sheetname].range(columnletter + str(wb.sheets[sheetname].cells.last_cell.row)).end('up').row + 1
        cell = columnletter + str(X)
        print(cell)
        print(cell[0], cell[1:])

        o = ord(str(cell[0]))
        col_num = o - 64
        row_num = int(cell[1:])

        wb.save(fp)
        wb.close()
        return (col_num, row_num)


    def modify_xl(dept2, coord):
        wb = load_workbook(filename=dept2)
        redFill = PatternFill(start_color='00FFCC99', end_color='00FFCC99', fill_type='solid')
        ws = wb["Sheet 1"]
        for row in ws.iter_rows():
            for cell in row:
                if (('11' in (cell.coordinate[1:3])) and len(cell.coordinate) == 3):
                    alignment = copy.copy(cell.alignment)
                    alignment.wrapText = True
                    alignment.horizontal = 'center'
                    alignment.vertical = 'bottom'
                    cell.alignment = alignment
                    ws[str(cell.coordinate)].font = Font(italic=True, bold=True)

        f_col = 'H10:H' + str(coord - 1)
        print(f_col)

        for row in ws.iter_rows():
            for cell in row:
                if (cell.coordinate[1:] not in ['1', '2', '3', '4', '5', '6', '7', '8', '9']):
                    alignment = copy.copy(cell.alignment)
                    alignment.wrapText = True
                    alignment.vertical = 'top'
                    cell.alignment = alignment

        for row in ws.iter_rows():
            for cell in row:
                if (cell.coordinate[0] == 'H' and int(cell.coordinate[1:3]) <= coord):
                    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                                         bottom=Side(style='thin'))
                    cell.border = thin_border
                    alignment = copy.copy(cell.alignment)
                    alignment.wrapText = True
                    alignment.vertical = 'center'
                    cell.alignment = alignment

        for cell in ws[f_col]:
            cell[0].fill = redFill  # light green

        wb.save(dept2)


    def merge_xl(dept2, d={'k1': 45, 'k2': 8, 'k3': 13, 'k4': 4, 'k5': 0, 'k6': 0},
                 d1={'co1': 18, 'co2': 32, 'co3': 0, 'co4': 0, 'co5': 0, 'co6': 0},
                 l1=[[3, 10, 11, ['CO', False], ['KL', True]], [3, 12, 13, ['CO', True], ['KL', False]]],
                 dist={"12": 50, "34": 50}, tot_marks=100):
        sheet.append([])

        plots(d, d1, l1, dist, tot_marks)


    from openpyxl.chart import *


    def plot_k(d):
        l1 = [k for k in d.keys()]
        l2 = [v for v in d.values()]
        l1.insert(0, "K Level ")
        l2.insert(0, "Marks")
        kvsm = list(zip(l1, l2))
        for val in kvsm:
            sheet.append(val)
        data = Reference(sheet, min_col=2, min_row=b + 2, max_row=b + len(l2) + 1)
        titles = Reference(sheet, min_col=1, min_row=b + 2, max_row=b + 1 + len(l1))
        chart = BarChart()
        chart.add_data(data=data)
        chart.set_categories(titles)
        chart.x_axis.title = " K level "
        chart.y_axis.title = " Marks "
        sheet.add_chart(chart, "E44")
        chart.title = " K level wise Marks Distribution"


    def plot_c(d):
        l1 = [k for k in d.keys()]
        l2 = [v for v in d.values()]
        l1.insert(0, "CO ")
        l2.insert(0, "Marks")
        cvsm = list(zip(l1, l2))
        for val in cvsm:
            sheet.append(val)
        data = Reference(sheet, min_col=2, min_row=9 + 1 + b + 1, max_row=9 + len(l2) + 1 + b)
        titles = Reference(sheet, min_col=1, min_row=10 + b + 1, max_row=10 + len(l1) + b)
        chart = BarChart()
        chart.add_data(data=data)
        chart.set_categories(titles)
        chart.x_axis.title = " CO level "
        chart.y_axis.title = " Marks "
        sheet.add_chart(chart, "H44")
        chart.title = " CO wise Marks Distribution"


    def oqn_chk(l1):
        # print("CHECK 0")
        sheet.append([])
        for n, i in enumerate(l1):
            l3 = []
            l4 = []
            if (i[3][1] == False):
                print("In part ", i[0], "\nQuestion numbers", i[1], i[2], "\nCO levels do not match")
                l3 = [("Q-No", "CO Match"), (i[1], False), (i[2], False)]
            if (i[4][1] == False):
                print("In part ", i[0], "\nQuestion numbers", i[1], i[2], "\nK levels do not match")
                l4 = [("Q-No", "K Match"), (i[1], False), (i[2], False)]
            if (len(l3) > 0):
                # print(l3)
                for val in l3:
                    # print("CHECK 1")
                    sheet.append(val)

            sheet.append([])
            if (len(l4) > 0):
                # print(l4)
                for val in l4:
                    # print("CHECK 2")
                    sheet.append(val)


    def plots(d={'k1': 45, 'k2': 8, 'k3': 13, 'k4': 4, 'k5': 0, 'k6': 0},
              d1={'co1': 18, 'co2': 32, 'co3': 0, 'co4': 0, 'co5': 0, 'co6': 0},
              l1=[[3, 10, 11, ['CO', False], ['KL', True]], [3, 12, 13, ['CO', True], ['KL', False]]],
              dist={"12": 50, "34": 50}, tot_marks=100):
        d2 = {}
        for k, v in d.items():
            k = k.upper()
            d2[k] = v
        plot_k(d2)
        sheet.append([])
        d3 = {}
        for k, v in d1.items():
            k = k.upper()
            d3[k] = v
        plot_c(d3)
        sheet.append([])
        sheet.append([])
        sheet.append([])
        sheet.append([])
        sheet.append([])
        sheet.append([])
        sheet.append([])
        sheet.append(["", "", "    %", "Marks"])
        sheet.append(["K1 / K2 Qns", "", round(dist["12"]), round(dist["12"] * tot_marks / 100)])  # dist['12']
        sheet.append(["K3 / K4 Qns", "", round(dist["34"]), round(dist["34"] * tot_marks / 100)])  # dist['34']
        sheet.append(["", "Total ", round(dist["12"]) + round(dist["34"]),
                      round(dist["12"] * tot_marks / 100) + round(dist["34"] * tot_marks / 100)])


    from datetime import datetime
    import os
    dept2 = "CSE.xlsx"


    def dept_dt_xl(dept2, sub_code):
        from datetime import datetime
        now = datetime.now()
        dt_string = now.strftime("%d/%m/%Y_%H:%M:%S")
        from openpyxl import Workbook
        wb = load_workbook(filename=dept2)
        ws = wb.active

        import datetime

        dt_string = dt_string.replace("/", "_")
        dt_string = dt_string.replace(":", "-")
        print(dt_string)

        dept2 = sub_code + "_" + dt_string + ".xlsx"

        wb.save(dept2)
        return dept2


    def comments_xl(Addcomments, a, b):
        sheet.append([])
        sheet.append([])
        c_after = sheet.cell(row=b, column=a)
        c_after.value = "                           NIL indicates the question does not use any action verb suggested in Bloom's taxonomy. "
        sheet.append(["Feedback for the faculty:"])
        sheet.append([])
        for n, i in enumerate(Addcomments):
            l1 = []
            s2 = ""
            l = [word.replace('\n', '') for word in i]
            l1.append(l)
            s2 = "".join(l)
            l2 = s2.split(".")

            for n2, j in enumerate(l2):
                tup1 = tuple(j.split("   "))
                tup2 = tuple(" ") + tup1

                sheet.append(tup2)


    a, b = end_coord()
    modify_xl(dept2, coord=b)
    sheet1 = load_workbook(filename=dept2)
    sheet = sheet1["Sheet 1"]
    merge_xl(dept2, Know_lev, Cour_Out, co_ch, dist, tot_marks)
    comments_xl(AddComments, a, b)
    sheet1.save(dept2)

    dept2 = dept_dt_xl(dept2, sub_code)
    os.remove("CSE.xls")
    os.remove("CSE.xlsx")
    os.remove("barChart.xlsx")
    os.startfile(dept2)
